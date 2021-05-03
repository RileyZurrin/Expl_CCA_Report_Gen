import csv
import docx
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import numpy as np
import pandas as pd


#Task Abbreviation
Task = "LDT"
#Groups
Groups = ['Controls', 'Bipolar Patients']

#Exploratory CCA Files. Enter data for all groups (in same order as line above above)
canoload = ['canoload_LDT_Con.xlsx', 'canoload_LDT_Bip.xlsx']
CCA = ['CCA_LDT_Con.csv', 'CCA_LDT_Bip.csv']
PCA = ['PCA_LDT_Con.csv', 'PCA_LDT_Bip.csv']

#Full Network Names from Classification (In order, C1 to Cn)
network = ['C1_Neg_91_TDMN_0.77', 'C2_Pos_86_LPN_1.52','C3_Pos_71_1RESP_1.14']
#Shortened Names (In same order as line above, with names matching the names in canoload.xlsx file)
networkshort = ['DMN', 'LPN', 'RESP']

#Open Variables.csv File
with open('Variables.csv') as csv_file:
    var = csv.reader(csv_file, delimiter=',')
    variables = list(var)[0]

#variables = ["Children"]

#Do Not Edit Below This Line
#--------------------------------------------------------------------------------------------------



#Network Matching Function (name extension)
def NMF(x):
    for n in range(len(network)):
        if x == networkshort[n]:
            return network[n]

#Dataset_Numbers of Networks (via canoload_TASK.xslx file)
df = pd.read_excel (canoload[0], sheet_name='Legend')

#Dataset_Number Matching Function
def DSNMF(x):
    for i in range(len(df)):
        if df.iloc[i,0] == x:
            return str(df.iloc[i,1])

mydoc = docx.Document()
#Document Style
# style = mydoc.styles['Normal']
# font = style.font
# font.name = 'Times New Roman'
# font.size = Pt(12)


#Header Style
style = mydoc.styles['Header']
font = style.font
font.size = Pt(12)
font.bold = True

#Read Conditions, Networks, and Variables
df2 = pd.read_excel (canoload[0], sheet_name='set_12_data_1')
conditions = []
constring = ''
for n in range(len(df2)):
    conditions.append(df2.iloc[n, 0])
conditionshort = []
for n in range(len(conditions)):
    conditionshort.append(conditions[n].split('_')[0])
conditionshort = list(dict.fromkeys(conditionshort))
for n in range(len(conditionshort)):
    if n != (len(conditionshort) - 1):
        constring += conditionshort[n] + ', '
    else:
        constring += 'and ' + conditionshort[n] + '.'
netstring = ''
for n in range(len(networkshort)):
    if n != len(networkshort) - 1:
        netstring += networkshort[n] + ', '
    else:
        netstring += 'and ' + networkshort[n] + '.'
varstring = ''
for n in range(len(variables)):
    if n != len(variables) - 1:
        varstring += variables[n] + ', '
    else:
        varstring += 'and ' + variables[n] + '.'
groupstring = ''
for n in range(len(Groups)):
    if n != len(Groups) - 1:
        groupstring += Groups[n] + ', '
    else:
        groupstring += 'and ' + Groups[n] + '.'
constringarray = ["Subjects in the " + conditionshort[j] + " condition were ___(explain " + conditionshort[j] + " condition)___. " for j in range(len(conditionshort))]
constringlong = ""
for n in range(len(constringarray)):
    constringlong += (constringarray[n])

# Write Report

mydoc.add_heading(Task + " Multiple-Set Canonical Correlation Report", 0)
mydoc.add_heading("Summary of Study")
if len(Groups) == 1:
    mydoc.add_paragraph(
        "This report is for the ________ (" + Task + ") task, in which participants were ___(describe task)___. There were " + str(
            len(conditionshort))
        + ' conditions: ' + constring + ' ' + constringlong)

if len(Groups) > 1:
    mydoc.add_paragraph(
        "This report is for the ________ (" + Task + ") task, in which participants were ___(describe task)___. There were " + str(
            len(Groups)) + " groups: " + groupstring + " There were " + str(
            len(conditionshort))
        + ' conditions: ' + constring + ' ' + constringlong)

mydoc.add_heading("Networks")
for l in range(len(network)):
    mydoc.add_paragraph(network[l])
mydoc.add_heading("Variables Included in Exploratory CCA")
mydoc.add_paragraph(' The increase to peak (ITP) and return to baseline (RTB) values for each network were included. '
'The behavioural variables used in the analysis are ' + varstring)

mydoc.add_heading("Notes")
mydoc.add_paragraph('1. Correlations (r values) were all made positive to ease interpretation, and '
                    'loadings were adjusted accordingly.' + '\n' + '    e.g. if r = -0.43 and loadings = '
                    '(0.32, 0.45, -0.11), then adjusted values would be r = 0.43'  + '\n' + '    and loadings = '
                    '(-0.32, -0.45 , 0.11).' + '\n' + '2. Loadings were bolded (equivalently, conditions were deemed significant) '
                                                     'according to the following formula:' + '\n'
                    + '    if abs(loading) >= abs(max(all_loadings))/2, then loading is bolded.' + '\n' + 'In words: a loading was bolded if its magnitude was '
                                                            'at least half as large as the loading with the largest magnitude '
                                                            'for the given interaction.')

for p in range(len(Groups)):
    if len(Groups) == 1:
        mydoc.add_heading("Summary of Results", 1)
        mydoc.add_heading("Summary for Each Behavioural Variable", 2)
        mydoc.add_paragraph("Only variables with significant results are reported.")
    else:
        mydoc.add_heading("Summary of Results for " + Groups[p], 1)
        mydoc.add_heading("Summary for Each Behavioural Variable", 2)
        mydoc.add_paragraph("Only variables with significant results are reported.")

    #Read CCA Data
    CCAmatrix = []
    sign = []
    with open(CCA[p]) as csv_file:
        data = csv.reader(csv_file, delimiter=',')
        rows = list(data)
        for j in range(0, len(variables)):
            matrix = []
            matrix.append([variables[j]])
            count = 1
            for q in range(0, len(networkshort)):
                for i in range(0, len(rows[0])):
                    if rows[4][i] == variables[j] and rows[2][i] == networkshort[q]:
                        neg = 1
                        interim = []
                        for k in [0,2,1,5]:
                            interim.append(np.transpose(rows)[i][k])
                        if float(interim[3]) < 0:
                            neg = -1
                            sign.append(-1)
                        else:
                            sign.append(1)
                        interim[3] = "%.2f" % abs(float(interim[3]))
                        matrix.append(interim)
                        num = []
                        for l in range(7,7 + len(conditions)):
                            num.append(neg*float(np.transpose(rows)[i][l]))
                        interim2 = ""
                        for m in range(0,len(conditions)):
                            if abs(num[m]) >= max([abs(ele) for ele in num])/2:
                                if np.sign(num[m]) == 1:
                                    interim2 = interim2 + rows[m+7][0] + ' positive, '
                                if np.sign(num[m]) == -1:
                                    interim2 = interim2 + rows[m + 7][0] + ' negative, '
                        matrix[count].append(interim2[:-2])
                        count += 1
            CCAmatrix.append(matrix)
            count += 1


    #read PCA data
    PCAmatrix = []
    with open(PCA[p]) as csv_file:
        data = csv.reader(csv_file, delimiter=',')
        rows = list(data)
        for j in range(0, len(variables)):
            matrix2 = []
            count = 1
            matrix2.append([variables[j]])
            for i in range(0, len(rows[0])):
                if rows[3][i] == variables[j]:
                    neg = 1
                    interim = []
                    for k in [0,1,4]:
                        interim.append(np.transpose(rows)[i][k])
                    if float(interim[2]) < 0:
                        neg = -1
                    interim[2] = "%.2f" % abs(float(interim[2]))
                    matrix2.append(interim)
                    num = []
                    for l in range(6,6 + len(conditions)):
                        num.append(neg*float(np.transpose(rows)[i][l]))
                    interim2 = ""
                    for m in range(0,len(conditions)):
                        if abs(num[m]) >= max([abs(ele) for ele in num])/2:
                            if np.sign(num[m]) == 1:
                                interim2 = interim2 + rows[m + 6][0] + ' positive, '
                            if np.sign(num[m]) == -1:
                                interim2 = interim2 + rows[m + 6][0] + ' negative, '
                    matrix2[count].append(interim2[:-2])
                    count += 1
            PCAmatrix.append(matrix2)
            count += 1

    #summary of results
    for i in range(len(variables)):
        insig = 0
        #PCA Results
        matrix2 = PCAmatrix[i]
        if len(matrix2) != 1:
            mydoc.add_heading("For " + variables[i] + ':', 3)
            table = mydoc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            row = table.rows[0]
            row.cells[0].text = 'PCA'
            col = table.columns[0]
            for cell in col.cells:
                cell.width = Cm(16.7)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = True
            table = mydoc.add_table(rows=len(matrix2), cols=4)
            table.style = 'Table Grid'
            # Column Widths
            col = table.columns[0]
            for cell in col.cells:
                cell.width = Cm(3.5)
            col2 = table.columns[1]
            for cell in col2.cells:
                cell.width = Cm(1.7)
            col3 = table.columns[2]
            for cell in col3.cells:
                cell.width = Cm(2)
            col4 = table.columns[3]
            for cell in col4.cells:
                cell.width = Cm(9)
            row = table.rows[0]
            row.cells[0].text = 'Network'
            row.cells[1].text = 'PC #'
            row.cells[2].text = 'r value'
            row.cells[3].text = 'Conditions'
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = True
            for j in range(1, len(matrix2)):
                row2 = table.rows[j]
                for k in range(4):
                    row2.cells[k].text = matrix2[j][k]
            mydoc.add_paragraph("")
        else:
            insig += 1
        #CCA Results
        matrix = CCAmatrix[i]
        if len(matrix) != 1:
            if insig == 1:
                mydoc.add_heading("For " + variables[i] + ':', 3)
                paragraph = mydoc.add_paragraph("No significant results with PCA.")
                paragraph.style = mydoc.styles['Header']
                paragraph = mydoc.add_paragraph("")
            table = mydoc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            row = table.rows[0]
            row.cells[0].text = 'CCA'
            col = table.columns[0]
            for cell in col.cells:
                cell.width = Cm(16.7)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = True
            table = mydoc.add_table(rows=len(matrix), cols=5)
            table.style = 'Table Grid'
            #Column Widths
            col = table.columns[0]
            for cell in col.cells:
                cell.width = Cm(3.4)
            col2 = table.columns[1]
            for cell in col2.cells:
                cell.width = Cm(2.5)
            col3 = table.columns[2]
            for cell in col3.cells:
                cell.width = Cm(1.2)
            col4 = table.columns[3]
            for cell in col4.cells:
                cell.width = Cm(2.5)
            col5 = table.columns[4]
            for cell in col5.cells:
                cell.width = Cm(6.6)
            row = table.rows[0]
            row.cells[0].text = 'Network Group'
            row.cells[1].text = 'Network'
            row.cells[2].text = 'CV #'
            row.cells[3].text = 'r value'
            row.cells[4].text = 'Conditions'
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = True
            for j in range(1, len(matrix)):
                row2 = table.rows[j]
                for k in range(5):
                    row2.cells[k].text = matrix[j][k]
        else:
            if insig == 1:
                pass
            else:
                paragraph = mydoc.add_paragraph("No significant results with CCA.")
                paragraph.style = mydoc.styles['Header']
    cols = np.transpose(rows)[:-1]
    mydoc.add_paragraph("")

    #Multiple-Set CCA and PCA Section
    mydoc.add_heading("Multiple-Set Canonical Correlation Analysis and Principal Component Analysis", 2)
    for i in range(len(variables)):
        matrix = CCAmatrix[i]
        if len(matrix) != 1:
            mydoc.add_heading("Significant Correlation with " + variables[i], 3)
        else:
            for n in range(len(cols)):
                if variables[i] in cols[n]:
                    mydoc.add_heading("Significant Correlation with " + variables[i], 3)
                    mydoc.add_paragraph("Significant correlation between " + cols[n][0] + ' and '
                                        + variables[i])
                    table = mydoc.add_table(rows=len(rows) + 1, cols=2)
                    table.style = 'Light List'
                    u = table.cell(0, 0)
                    v = table.cell(0, 1)
                    u.merge(v)
                    u.text = "PCA"
                    col = table.columns[0]
                    for a in range(len(rows)):
                        col.cells[a + 1].text = cols[0][a]
                    col.cells[5].text = 'r'
                    col2 = table.columns[1]
                    cond = []
                    for b in range(6, 6 + len(conditions)):
                        cond.append(abs(float(cols[n][b])))
                    sgn = np.sign(float(cols[n][4]))
                    for k in range(len(rows)):
                        try:
                            if round(float(cols[n][k]), 2) < 1:
                                col2.cells[k + 1].text = "%.2f" % (sgn * float(cols[n][k]))
                            else:
                                col2.cells[k + 1].text = str(int(cols[n][k]))
                        except ValueError:
                            col2.cells[k + 1].text = cols[n][k]
                    t = 2
                    while round(float(cols[n][5]), t) == 0:
                        t += 1
                    col2.cells[6].text = str(round(float(cols[n][5]), t))
                    for cell in col2.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                try:
                                    if 1 > abs(float(cell.text)) >= max(cond) / 2:
                                        font = run.font
                                        font.bold = True
                                except ValueError:
                                    pass
                    # Column Widths
                    col = table.columns[0]
                    for cell in col.cells:
                        cell.width = Cm(4)
                    col2 = table.columns[1]
                    for cell in col2.cells:
                        cell.width = Cm(5)
                    mydoc.add_paragraph("")
        for j in range(1, len(matrix)):
            coldel = []
            for n in range(len(cols)):
                if variables[i] in cols[n] and matrix[j][1] in cols[n]:
                    mydoc.add_paragraph("Significant correlation between " + matrix[j][1] + ' and '
                                        + variables[i])
                    table = mydoc.add_table(rows=len(rows) + 1, cols=2)
                    table.style = 'Light List'
                    u = table.cell(0, 0)
                    v = table.cell(0, 1)
                    u.merge(v)
                    u.text = "PCA"
                    col = table.columns[0]
                    for a in range(len(rows)):
                        col.cells[a + 1].text = cols[0][a]
                    col.cells[5].text = 'r'
                    col2 = table.columns[1]
                    cond = []
                    for b in range(6, 6 + len(conditions)):
                        cond.append(abs(float(cols[n][b])))
                    sgn = np.sign(float(cols[n][4]))
                    for k in range(len(rows)):
                        try:
                            if round(float(cols[n][k]), 2) < 1:
                                col2.cells[k + 1].text = "%.2f" % (sgn * float(cols[n][k]))
                            else:
                                col2.cells[k + 1].text = str(int(cols[n][k]))
                        except ValueError:
                            col2.cells[k + 1].text = cols[n][k]
                    t = 2
                    while round(float(cols[n][5]), t) == 0:
                        t += 1
                    col2.cells[6].text = str(round(float(cols[n][5]), t))
                    for cell in col2.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            for run in paragraph.runs:
                                try:
                                    if 1 > abs(float(cell.text)) >= max(cond) / 2:
                                        font = run.font
                                        font.bold = True
                                except ValueError:
                                    pass
                    # Column Widths
                    col = table.columns[0]
                    for cell in col.cells:
                        cell.width = Cm(4)
                    col2 = table.columns[1]
                    for cell in col2.cells:
                        cell.width = Cm(5)
                    mydoc.add_paragraph("")
                    coldel.append(n)
            cols = np.delete(cols, coldel, 0)
            CV = matrix[j][2]
            Networks = matrix[j][0].split('&')
            Head4 = ''
            sheet = ''
            for k in range(len(Networks)):
                NMF(Networks[k])
                try:
                    Head4 = Head4 + NMF(Networks[k]) + ' & ' + '\n'
                except TypeError:
                    print("Error: The names you inputted into networkshort (e.g. DMN_AAR) do not match the names in the canoload Excel file.")
                    exit()
                sheet = sheet + DSNMF(Networks[k])
            Networks.remove(matrix[j][1])
            Networks.insert(0, matrix[j][1])
            Head4 = Head4[:-3]
            mydoc.add_heading(Head4, 4)
            mydoc.add_paragraph("Significant correlation between " + matrix[j][1] + ' and '
                            + variables[i])
            table = mydoc.add_table(rows = 6 + len(conditions), cols = 1 + len(Networks))
            a = table.cell(1,1)
            b = table.cell(1, len(Networks))
            a.merge(b)
            NetCombo = ''
            for m in range(len(Networks)):
                NetCombo = NetCombo + Networks[m] + ' & '
            a.text = NetCombo[:-3]
            a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            col = table.columns[0]
            col.cells[0].text = 'CCA'
            col.cells[1].text = 'Network Group'
            col.cells[2].text = 'CV #'
            col.cells[3].text = 'r'
            col.cells[4].text = 'Network'
            col.cells[5].text = 'Condition'
            for s in range(len(conditions)):
                col.cells[6 + s].text = conditions[s]
            row2 = table.rows[4]
            for s in range(len(Networks)):
                row2.cells[s + 1].text = Networks[s]
            col2 = table.columns[1]
            col2.cells[2].text = CV
            col2.cells[3].text = matrix[j][3]
            a = table.cell(5, 1)
            b = table.cell(5, len(Networks))
            a.merge(b)
            a.text = 'Loadings'
            a.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for l in range(len(Networks)):
                maxm = []
                col = table.columns[l + 1]
                df1 = pd.read_excel (canoload[p], sheet_name='set_' + sheet + '_data_' + DSNMF(Networks[l]))
                for m in range(len(conditions)):
                    col.cells[6 + m].text = "%.2f" % (sign[j - 1] * float(df1.iloc[m, int(CV)]))
                    maxm.append(abs(round(df1.iloc[m,int(CV)],2)))
                count = 0
                for cell in col.cells:
                    paragraphs = cell.paragraphs
                    paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            if l == 0 and count > 4:
                                if abs(float(cell.text)) >= max(maxm)/2:
                                    font = run.font
                                    font.bold = True
                            if l == 0 and count == 4:
                                font = run.font
                                font.bold = True
                            if l > 0 and count > 2:
                                if abs(float(cell.text)) >= max(maxm)/2:
                                    font = run.font
                                    font.bold = True
                            else:
                                pass
                            count += 1
            for row in table.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Cm(0.5)
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
            table.style = 'Light List Accent 1'
            mydoc.add_paragraph("")
mydoc.save(Task + "_Multiple_Set_CCA_Report.docx")
