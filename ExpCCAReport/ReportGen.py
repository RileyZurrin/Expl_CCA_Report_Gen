import csv
import docx
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Pt
import numpy as np
import pandas as pd


#Task Abbreviation
Task = "RSPM"

#Upload Files
canoload = 'canoload_RSPM_con.xlsx'
CCA = 'CCA_RSPM_con.csv'
PCA = 'PCA_RSPM_con.csv'

#Open Variables.txt File
with open('Variables.txt') as csv_file:
    var = csv.reader(csv_file, delimiter=',')
    variables = list(var)[0]

#Network Names
network = ['C1_Neg_92_TDMN_0.69 / PP_NegL92_AAR_0.57', 'C2_Pos_92_1RESP_0.78','C3_Pos_78_EXT_1.20',
           'C4_Pos_63_INIT_0.91 / Neg_37_TDMN_0.80','C5_Pos_88_CE_0.78 / PP_NegL_12_1RESP_0.72']
#Shortened Names
networkshort = ['DMN_AAR', 'RESP', 'EXT', 'INIT_DMN', 'CE_RESP']

#Network Matching Function (name extension)
def NMF(x):
    for n in range(len(network)):
        if x == networkshort[n]:
            return network[n]

#Dataset_Numbers of Networks (via canoload_TASK.xslx file)
df = pd.read_excel (canoload, sheet_name='Legend')

#Dataset_Number Matching Function
def DSNMF(x):
    for i in range(len(df)):
        if df.iloc[i,0] == x:
            return str(df.iloc[i,1])

mydoc = docx.Document()
#Document Style
# style = mydoc.styles['Normal']
# font = style.font
# font.name = 'Times New Roman'
# font.size = Pt(12)

#Header Style
style = mydoc.styles['Header']
font = style.font
font.size = Pt(12)
font.bold = True

#Read Conditions, Networks, and Variables
df2 = pd.read_excel (canoload, sheet_name='set_12_data_1')
conditions = []
constring = ''
for n in range(len(df2)):
    conditions.append(df2.iloc[n, 0])
    if n != len(df2) - 1:
        constring += conditions[n] + ', '
    else:
        constring += 'and ' + conditions[n] + '.'
netstring = ''
for n in range(len(networkshort)):
    if n != len(networkshort) - 1:
        netstring += networkshort[n] + ', '
    else:
        netstring += 'and ' + networkshort[n] + '.'
varstring = ''
for n in range(len(variables)):
    if n != len(variables) - 1:
        varstring += variables[n] + ', '
    else:
        varstring += 'and ' + variables[n] + '.'

#Read CCA Data
CCAmatrix = []
with open(CCA) as csv_file: #open CCA_TASK.csv file
    data = csv.reader(csv_file, delimiter=',')
    rows = list(data)
    for j in range(0, len(variables)):
        matrix = []
        matrix.append([variables[j]])
        count = 1
        for i in range(0, len(rows[0])):
            if rows[4][i] == variables[j]:
                interim = []
                for k in [0,2,1,5]:
                    interim.append(np.transpose(rows)[i][k])
                interim[3] = str(round(float(interim[3]), 2))
                matrix.append(interim)
                num = []
                for l in range(7,7 + len(conditions)):
                    num.append(float(np.transpose(rows)[i][l]))
                interim2 = ""
                for m in range(0,len(conditions)):
                    if abs(num[m]) >= max([abs(ele) for ele in num])/2:
                        if np.sign(num[m]) == 1:
                            interim2 = interim2 + rows[m+7][0] + ' positive, '
                        if np.sign(num[m]) == -1:
                            interim2 = interim2 + rows[m + 7][0] + ' negative, '
                matrix[count].append(interim2[:-2])
                count += 1
        CCAmatrix.append(matrix)
        count += 1

PCAmatrix = []
with open(PCA) as csv_file: #open PCA_TASK.csv file
    data = csv.reader(csv_file, delimiter=',')
    rows = list(data)
    for j in range(0, len(variables)):
        matrix2 = []
        count = 1
        matrix2.append([variables[j]])
        for i in range(0, len(rows[0])):
            if rows[3][i] == variables[j]:
                interim = []
                for k in [0,1,4]:
                    interim.append(np.transpose(rows)[i][k])
                interim[2] = str(round(float(interim[2]), 2))
                matrix2.append(interim)
                num = []
                for l in range(6,6 + len(conditions)):
                    num.append(float(np.transpose(rows)[i][l]))
                interim2 = ""
                for m in range(0,len(conditions)):
                    if abs(num[m]) >= max([abs(ele) for ele in num])/2:
                        if np.sign(num[m]) == 1:
                            interim2 = interim2 + rows[m + 6][0] + ' positive, '
                        if np.sign(num[m]) == -1:
                            interim2 = interim2 + rows[m + 6][0] + ' negative, '
                matrix2[count].append(interim2[:-2])
                count += 1
        PCAmatrix.append(matrix2)
        count += 1

#Write Report

mydoc.add_heading(Task + " Multiple-Set Canonical Correlation Report", 0)
mydoc.add_heading("Summary of Study, Networks, Conditions, and Behavioural Data")
mydoc.add_paragraph("This report is for the ________ (" + Task + ") task, in which participants were _____. There were " + str(len(conditions))
                    + ' conditions: ' + constring + ' Subjects in the _____ condition did ______, subjects in the ____ conditions '
                                                 'did _____, etc. The ' + str(len(network)) + ' networks identified as part of this study are: '
                    + netstring + ' Variables inputted include the ITP and RTB values for each network. '
                                  'Behavioural variables used in the analysis are ' + varstring)
mydoc.add_heading("Summary of Results",1)
mydoc.add_heading("Summary for Each Behavioural Variable",2)

#summary of results
for i in range(len(variables)):
    mydoc.add_heading("For " + variables[i] + ':', 3)
    #PCA Results
    matrix2 = PCAmatrix[i]
    if len(matrix2) != 1:
        table = mydoc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        row = table.rows[0]
        row.cells[0].text = 'PCA'
        col = table.columns[0]
        for cell in col.cells:
            cell.width = Cm(16.7)
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.bold = True
        table = mydoc.add_table(rows=len(matrix2), cols=4)
        table.style = 'Table Grid'
        # Column Widths
        col = table.columns[0]
        for cell in col.cells:
            cell.width = Cm(5)
        col2 = table.columns[1]
        for cell in col2.cells:
            cell.width = Cm(1.7)
        col3 = table.columns[2]
        for cell in col3.cells:
            cell.width = Cm(2)
        col4 = table.columns[3]
        for cell in col4.cells:
            cell.width = Cm(7.5)
        row = table.rows[0]
        row.cells[0].text = 'Network'
        row.cells[1].text = 'PC #'
        row.cells[2].text = 'R value'
        row.cells[3].text = 'Conditions'
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.bold = True
        for j in range(1, len(matrix2)):
            row2 = table.rows[j]
            for k in range(4):
                row2.cells[k].text = matrix2[j][k]
        mydoc.add_paragraph("")
    else:
        paragraph = mydoc.add_paragraph("No significant results with PCA.")
        paragraph.style = mydoc.styles['Header']
    #CCA Results
    matrix = CCAmatrix[i]
    if len(matrix) != 1:
        table = mydoc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        row = table.rows[0]
        row.cells[0].text = 'CCA'
        col = table.columns[0]
        for cell in col.cells:
            cell.width = Cm(16.7)
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.bold = True
        table = mydoc.add_table(rows=len(matrix), cols=5)
        table.style = 'Table Grid'
        #Column Widths
        col = table.columns[0]
        for cell in col.cells:
            cell.width = Cm(5)
        col2 = table.columns[1]
        for cell in col2.cells:
            cell.width = Cm(2.5)
        col3 = table.columns[2]
        for cell in col3.cells:
            cell.width = Cm(1.2)
        col4 = table.columns[3]
        for cell in col4.cells:
            cell.width = Cm(2.5)
        col5 = table.columns[4]
        for cell in col5.cells:
            cell.width = Cm(5)
        row = table.rows[0]
        row.cells[0].text = 'Network Group'
        row.cells[1].text = 'Network'
        row.cells[2].text = 'CV #'
        row.cells[3].text = 'R value'
        row.cells[4].text = 'Conditions'
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.bold = True
        for j in range(1, len(matrix)):
            row2 = table.rows[j]
            for k in range(5):
                row2.cells[k].text = matrix[j][k]
    else:
        paragraph = mydoc.add_paragraph("No significant results with CCA.")
        paragraph.style = mydoc.styles['Header']

cols = np.transpose(rows)[:-1]
#Principle Component Analysis Section
mydoc.add_heading("Principal Component Analysis", 2)
for i in range(len(cols) - 1):
    table = mydoc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    col = table.columns[0]
    for j in range(len(rows)):
        col.cells[j].text = cols[0][j]
    col2 = table.columns[1]
    cond = []
    for l in range(6, 6 + len(conditions)):
        cond.append(abs(float(cols[i+1][l])))
    for k in range(len(rows)):
        try:
            col2.cells[k].text = str(round(float(cols[i + 1][k]),2))
        except ValueError:
            col2.cells[k].text = cols[i+1][k]
    for cell in col2.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                try:
                    if abs(float(cell.text)) >= max(cond) / 2:
                        font = run.font
                        font.bold = True
                except ValueError:
                    pass
    # Column Widths
    col = table.columns[0]
    for cell in col.cells:
        cell.width = Cm(4)
    col2 = table.columns[1]
    for cell in col2.cells:
        cell.width = Cm(5)
    mydoc.add_paragraph("")

#Multiple-Set Canonical Correlation Analysis Section
mydoc.add_heading("Multiple-Set Canonical Correlation Analysis", 2)
for i in range(len(variables)):
    matrix = CCAmatrix[i]
    if len(matrix) != 1:
        mydoc.add_heading("Significant Correlation with " + variables[i], 3)
    for j in range(1, len(matrix)):
        CV = matrix[j][2]
        Networks = matrix[j][0].split('&')
        Head4 = ''
        sheet = ''
        for k in range(len(Networks)):
            NMF(Networks[k])
            Head4 = Head4 + NMF(Networks[k]) + ' & '
            sheet = sheet + DSNMF(Networks[k])
        Networks.remove(matrix[j][1])
        Networks.insert(0, matrix[j][1])
        Head4 = Head4[:-2]
        mydoc.add_heading(Head4, 4)
        mydoc.add_paragraph("Significant correlation between " + matrix[j][1] + ' and '
                            + variables[i] + '\n' + "Canonical Variate " + CV
                            + " (R = " + matrix[j][3] + ", P < 0.01)")
        for l in range(len(Networks)):
            maxm = []
            if Networks[l] == matrix[j][1]:
                paragraph = mydoc.add_paragraph(Networks[l] + ' CV' + CV + " Canonical Loadings")
                paragraph.style = mydoc.styles['Header']
            else:
                paragraph = mydoc.add_paragraph(Networks[l] + ' CV' + CV + " Canonical Loadings - Not Significant")
                paragraph.style = mydoc.styles['Header']
            df1 = pd.read_excel (canoload, sheet_name='set_' + sheet + '_data_' + DSNMF(Networks[l]))
            table = mydoc.add_table(rows=1 + len(conditions), cols=2)
            col = table.columns[0]
            col.cells[0].text = 'Condition'
            for n in range(1,1 + len(conditions)):
                col.cells[n].text = df1.iloc[n-1,0]
            col2 = table.columns[1]
            for m in range(1,1 + len(conditions)):
                col2.cells[m].text = str(round(df1.iloc[m - 1,int(CV)],2))
                maxm.append(abs(round(df1.iloc[m - 1,int(CV)],2)))
            for cell in col2.cells:
                cell.width = Cm(2)
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        if abs(float(cell.text)) >= max(maxm)/2:
                            font = run.font
                            font.bold = True
            for row in table.rows:
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row.height = Cm(0.5)
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(10)
mydoc.save(Task + "_Multiple_Set_CCA_Report.docx")
